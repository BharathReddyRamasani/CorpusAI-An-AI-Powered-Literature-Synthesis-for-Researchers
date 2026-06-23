"""
Services — Report Generation Service
Collects metadata + summary + citations + QA insights
→ Gemini generates report content
→ Exports PDF and DOCX
→ Stores in SQLite
"""

import asyncio
import logging
from pathlib import Path

import httpx
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import settings
from app.models.citation import Citation
from app.models.chat import ChatHistory
from app.models.paper import Paper
from app.models.report import Report
from app.models.summary import Summary
from app.prompts.report import REPORT_SYSTEM_PROMPT, REPORT_USER_PROMPT
from app.utils.exceptions import BadRequestException, NotFoundException, ServiceException
from app.utils.helpers import get_report_path

from app.utils.groq_client import call_groq_api_with_rotation

logger = logging.getLogger("app")


# ── Report Content Generation ─────────────────────────────────────────────────

async def _generate_report_content(db: AsyncSession, paper_id: str) -> str:
    """Collect all data and generate report content via Gemini."""
    # Get paper
    paper_result = await db.execute(select(Paper).where(Paper.paper_id == paper_id))
    paper = paper_result.scalar_one_or_none()
    if not paper:
        raise NotFoundException("Paper")

    if paper.status != "ready":
        raise BadRequestException(
            f"Paper is not ready for report generation (status='{paper.status}')."
        )

    # Get summary
    summary_result = await db.execute(
        select(Summary).where(Summary.paper_id == paper_id)
    )
    summary_obj = summary_result.scalars().first()
    summary_text = summary_obj.summary if summary_obj else "No summary available."

    # Get citations (limit to 10 for Groq free tier limit)
    citations_result = await db.execute(
        select(Citation).where(Citation.paper_id == paper_id).limit(10)
    )
    citations = list(citations_result.scalars().all())
    citations_text = "\n".join(
        f"- {c.author or 'Unknown'} ({c.year or 'n.d.'}) — {c.title or c.raw_text or ''}"
        for c in citations
    ) or "No citations extracted."

    # Get QA insights (limit to 5)
    qa_result = await db.execute(
        select(ChatHistory).where(ChatHistory.paper_id == paper_id).limit(5)
    )
    qa_history = list(qa_result.scalars().all())
    qa_text = "\n".join(
        f"Q: {qa.question}\nA: {qa.answer}" for qa in qa_history
    ) or "No Q&A interactions recorded."

    # Build prompt (truncate fields for 8000 TPM limit)
    prompt = REPORT_USER_PROMPT.format(
        title=(paper.title or paper.filename)[:500],
        authors=(paper.authors or "Unknown")[:500],
        abstract=(paper.abstract or "Not available")[:2000],
        summary=summary_text[:4000],
        citation_count=len(citations),
        citations=citations_text[:2000],
        qa_insights=qa_text[:2000],
    )

    logger.info(f"Generating report content via Groq for paper_id={paper_id}")

    try:
        return await call_groq_api_with_rotation(prompt, REPORT_SYSTEM_PROMPT, max_tokens=2000)
    except Exception as e:
        logger.error(f"Groq report generation failed: {e}", exc_info=True)
        raise ServiceException(f"Report generation failed: {str(e)}")


# ── PDF Export ────────────────────────────────────────────────────────────────

def _export_pdf(content: str, output_path: Path, title: str) -> None:
    """Export report content to a PDF file using ReportLab."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=2.5 * cm,
        rightMargin=2.5 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ReportTitle",
        parent=styles["Title"],
        fontSize=18,
        spaceAfter=12,
        textColor=colors.HexColor("#1a237e"),
    )
    h1_style = ParagraphStyle(
        "H1",
        parent=styles["Heading1"],
        fontSize=14,
        spaceBefore=16,
        spaceAfter=8,
        textColor=colors.HexColor("#283593"),
    )
    body_style = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontSize=10,
        spaceAfter=8,
        leading=14,
    )

    story = []
    story.append(Paragraph(f"Research Analysis Report", title_style))
    story.append(Paragraph(f"<i>{title}</i>", styles["Heading2"]))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#283593")))
    story.append(Spacer(1, 0.5 * cm))

    # Parse markdown-style headings
    for line in content.split("\n"):
        line = line.strip()
        if not line:
            story.append(Spacer(1, 0.2 * cm))
        elif line.startswith("## "):
            story.append(Paragraph(line[3:], h1_style))
        elif line.startswith("# "):
            story.append(Paragraph(line[2:], title_style))
        elif line.startswith("- ") or line.startswith("* "):
            story.append(Paragraph(f"• {line[2:]}", body_style))
        else:
            # Escape special chars for reportlab
            safe_line = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(safe_line, body_style))

    doc.build(story)
    logger.info(f"PDF report exported: {output_path}")

def _export_pdf_stream(content: str, title: str) -> bytes:
    """Export report content to a PDF in memory."""
    import io
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=2.5 * cm,
        rightMargin=2.5 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ReportTitle",
        parent=styles["Title"],
        fontSize=18,
        spaceAfter=12,
        textColor=colors.HexColor("#1a237e"),
    )
    h1_style = ParagraphStyle(
        "H1",
        parent=styles["Heading1"],
        fontSize=14,
        spaceBefore=16,
        spaceAfter=8,
        textColor=colors.HexColor("#283593"),
    )
    body_style = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontSize=10,
        spaceAfter=8,
        leading=14,
    )

    story = []
    story.append(Paragraph(f"Research Analysis Report", title_style))
    story.append(Paragraph(f"<i>{title}</i>", styles["Heading2"]))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#283593")))
    story.append(Spacer(1, 0.5 * cm))

    for line in content.split("\n"):
        line = line.strip()
        if not line:
            story.append(Spacer(1, 0.2 * cm))
        elif line.startswith("## "):
            story.append(Paragraph(line[3:], h1_style))
        elif line.startswith("# "):
            story.append(Paragraph(line[2:], title_style))
        elif line.startswith("- ") or line.startswith("* "):
            story.append(Paragraph(f"• {line[2:]}", body_style))
        else:
            safe_line = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(safe_line, body_style))

    doc.build(story)
    pdf_bytes = buffer.getvalue()
    buffer.close()
    return pdf_bytes


# ── DOCX Export ───────────────────────────────────────────────────────────────

def _export_docx(content: str, output_path: Path, title: str) -> None:
    """Export report content to a DOCX file using python-docx."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    # Title
    title_para = doc.add_heading(f"Research Analysis Report", 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading(title, level=1)
    doc.add_paragraph()

    # Parse content
    for line in content.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph()
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)

    doc.save(str(output_path))
    logger.info(f"DOCX report exported: {output_path}")

def _export_docx_stream(content: str, title: str) -> bytes:
    """Export report content to a DOCX in memory."""
    import io
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    buffer = io.BytesIO()
    doc = Document()

    title_para = doc.add_heading(f"Research Analysis Report", 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading(title, level=1)
    doc.add_paragraph()

    for line in content.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph()
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)

    doc.save(buffer)
    docx_bytes = buffer.getvalue()
    buffer.close()
    return docx_bytes


# ── Main Report Generation Function ──────────────────────────────────────────

async def generate_report(
    db: AsyncSession,
    paper_id: str,
    fmt: str = "pdf",
) -> Report:
    """
    Generate a PDF or DOCX research report for the given paper.

    Args:
        db: Async database session.
        paper_id: Target paper ID.
        fmt: Export format — 'pdf' or 'docx'.

    Returns:
        Report ORM object.
    """
    if fmt not in ("pdf", "docx"):
        raise BadRequestException("Format must be 'pdf' or 'docx'.")

    # Generate content via Gemini
    content = await _generate_report_content(db, paper_id)

    # Get paper title for file header
    paper_result = await db.execute(select(Paper).where(Paper.paper_id == paper_id))
    paper = paper_result.scalar_one_or_none()
    paper_title = paper.title or paper.filename if paper else paper_id

    # Export to file
    output_path = get_report_path(paper_id, fmt)

    loop = asyncio.get_event_loop()
    if fmt == "pdf":
        await loop.run_in_executor(None, _export_pdf, content, output_path, paper_title)
    else:
        await loop.run_in_executor(None, _export_docx, content, output_path, paper_title)

    # Store report record
    report = Report(
        paper_id=paper_id,
        report_path=str(output_path),
        format=fmt,
    )
    db.add(report)
    await db.flush()
    await db.refresh(report)
    logger.info(f"Report record stored: id={report.id} format={fmt}")
    return report

async def generate_report_stream(
    db: AsyncSession,
    paper_id: str,
    fmt: str = "pdf",
) -> bytes:
    """
    Generate a PDF or DOCX research report entirely in memory.
    Returns the raw bytes for streaming to the client.
    """
    if fmt not in ("pdf", "docx"):
        raise BadRequestException("Format must be 'pdf' or 'docx'.")

    # Generate content via Gemini
    content = await _generate_report_content(db, paper_id)

    # Get paper title for file header
    paper_result = await db.execute(select(Paper).where(Paper.paper_id == paper_id))
    paper = paper_result.scalar_one_or_none()
    paper_title = paper.title or paper.filename if paper else paper_id

    loop = asyncio.get_event_loop()
    if fmt == "pdf":
        file_bytes = await loop.run_in_executor(None, _export_pdf_stream, content, paper_title)
    else:
        file_bytes = await loop.run_in_executor(None, _export_docx_stream, content, paper_title)

    # Store report record even for streaming so the dashboard count updates
    from app.models.report import Report
    report = Report(
        paper_id=paper_id,
        report_path="streamed_in_memory",
        format=fmt,
    )
    db.add(report)
    await db.flush()

    logger.info(f"Dynamic report generated in memory format={fmt}")
    return file_bytes
